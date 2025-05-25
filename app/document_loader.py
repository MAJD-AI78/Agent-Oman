import os
import mammoth
import shutil
import pytesseract
import pandas as pd
from PIL import Image
from docx import Document
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader
)
from pptx import Presentation

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".pptx", ".xlsx", ".csv", ".png", ".jpg", ".jpeg", ".txt", ".md"}
SKIPPED_FILES = []
LOADED_FILES = []

DOCUMENTS_DIR = "documents"


def convert_doc_to_docx(filepath):
    docx_path = filepath + "x"
    with open(filepath, "rb") as f:
        result = mammoth.convert_to_html(f)
        doc = Document()
        doc.add_paragraph(result.value)
        doc.save(docx_path)
    return docx_path


def load_pdf(path):
    return PyPDFLoader(path).load()


def load_docx(path):
    return UnstructuredWordDocumentLoader(path).load()


def load_doc(path):
    try:
        docx_path = convert_doc_to_docx(path)
        return load_docx(docx_path)
    except Exception as e:
        SKIPPED_FILES.append(f"{path} (conversion error: {e})")
        return []


def load_excel(path):
    try:
        df = pd.read_excel(path, sheet_name=None)
        text = "\n".join(
            df[sheet].to_csv(index=False) for sheet in df
        )
        return [Document(page_content=text)]
    except Exception as e:
        SKIPPED_FILES.append(f"{path} (Excel load error: {e})")
        return []


def load_csv(path):
    try:
        df = pd.read_csv(path)
        return [Document(page_content=df.to_csv(index=False))]
    except Exception as e:
        SKIPPED_FILES.append(f"{path} (CSV load error: {e})")
        return []


def load_image(path):
    try:
        text = pytesseract.image_to_string(Image.open(path))
        return [Document(page_content=text)]
    except Exception as e:
        SKIPPED_FILES.append(f"{path} (OCR error: {e})")
        return []


def load_text(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return [Document(page_content=f.read())]
    except Exception as e:
        SKIPPED_FILES.append(f"{path} (Text load error: {e})")
        return []


def clean_document_folder():
    print("🧼 Cleaning folders...")
    for root, dirs, files in os.walk(DOCUMENTS_DIR):
        for folder in dirs:
            folder_path = os.path.join(root, folder)
            if not folder_path.lower().endswith((".pptx_files",)):
                print(f"❌ Removing folder: {folder_path}")
                shutil.rmtree(folder_path, ignore_errors=True)


def load_documents():
    print("📥 Starting document loading...")
    all_docs = []
    clean_document_folder()

    for root, _, files in os.walk(DOCUMENTS_DIR):
        for file in files:
            ext = os.path.splitext(file)[-1].lower()
            path = os.path.join(root, file)

            if ext not in SUPPORTED_EXTENSIONS:
                SKIPPED_FILES.append(f"{file} (unsupported extension)")
                continue

            try:
                if ext == ".pdf":
                    docs = load_pdf(path)
                elif ext == ".docx":
                    docs = load_docx(path)
                elif ext == ".doc":
                    docs = load_doc(path)
                elif ext in [".xlsx"]:
                    docs = load_excel(path)
                elif ext == ".csv":
                    docs = load_csv(path)
                elif ext in [".png", ".jpg", ".jpeg"]:
                    docs = load_image(path)
                elif ext in [".txt", ".md"]:
                    docs = load_text(path)
                else:
                    SKIPPED_FILES.append(f"{file} (not handled)")
                    continue

                if docs:
                    all_docs.extend(docs)
                    LOADED_FILES.append(file)

            except Exception as e:
                SKIPPED_FILES.append(f"{file} (load error: {e})")

    print(f"✅ Loaded {len(LOADED_FILES)} files.")
    print("📋 Loaded:", LOADED_FILES)
    print("⚠️ Skipped:", SKIPPED_FILES)

    if not all_docs:
        print("⚠️ No documents were loaded.")
        return

    # Save to vectorstore
    print("💾 Saving to FAISS vectorstore...")
    embeddings = OpenAIEmbeddings()
    db = FAISS.from_documents(all_docs, embeddings)
    db.save_local("vectorstore")
    print("✅ Vectorstore saved.")


if __name__ == "__main__":
    load_documents()